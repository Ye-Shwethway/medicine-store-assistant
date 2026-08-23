from __future__ import annotations

import io
import json
import re
from datetime import datetime, timezone
from typing import Any, Literal

from docx import Document
from fastapi import APIRouter, Depends, HTTPException, Query, Response
from sqlalchemy import text

from app.ai_workspace_access import require_ai_chat_access
from app.ai_workspace_chat import _message_attachment_map, _owned_conversation
from app.dashboard_auth import _engine, require_owner_session
from app.multi_agent_review import _work_item_detail

router = APIRouter(prefix="/dashboard/api/ai-workspace", tags=["conversation-export"])

SNAPSHOT_SCHEMA_VERSION = "2026-08-24.v1"
ExportFormat = Literal["docx", "json"]


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _filename(value: str, suffix: str) -> str:
    base = re.sub(r"[^A-Za-z0-9._-]+", "-", value.strip()).strip("-.") or "msa-snapshot"
    return f"{base[:90]}.{suffix}"


def _download(content: bytes, *, media_type: str, filename: str) -> Response:
    return Response(
        content=content,
        media_type=media_type,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Cache-Control": "no-store",
            "Pragma": "no-cache",
            "X-Content-Type-Options": "nosniff",
        },
    )


def _json_bytes(snapshot: dict[str, Any]) -> bytes:
    return json.dumps(snapshot, ensure_ascii=False, indent=2, default=str).encode("utf-8")


def _clean_human_text(value: Any) -> str:
    text_value = str(value or "")
    text_value = re.sub(r"^\s*#{1,6}\s+", "", text_value, flags=re.MULTILINE)
    text_value = re.sub(r"\*\*([^*]+)\*\*", r"\1", text_value)
    text_value = re.sub(r"__([^_]+)__", r"\1", text_value)
    text_value = re.sub(r"`([^`\n]+)`", r"\1", text_value)
    text_value = re.sub(r"^\s*\|?\s*:?-{3,}.*\|?\s*$", "", text_value, flags=re.MULTILINE)
    lines = []
    for line in text_value.splitlines():
        stripped = line.strip()
        if stripped == "---":
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            line = " · ".join(cell for cell in cells if cell)
        lines.append(line.rstrip())
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()


def _chat_snapshot(connection: Any, conversation_id: str, principal: dict[str, str]) -> dict[str, Any]:
    conversation = _owned_conversation(connection, conversation_id, principal)
    rows = connection.execute(
        text(
            """
            SELECT message_id::text AS message_id, role, content, runtime_provenance, created_at
            FROM ai_workspace_messages
            WHERE conversation_id=CAST(:conversation_id AS uuid)
            ORDER BY created_at, CASE role WHEN 'USER' THEN 0 ELSE 1 END, message_id
            """
        ),
        {"conversation_id": conversation_id},
    ).mappings().all()
    attachments = _message_attachment_map(connection, conversation_id, principal)
    messages: list[dict[str, Any]] = []
    for row in rows:
        item = dict(row)
        item["attachments"] = attachments.get(item["message_id"], [])
        messages.append(item)
    return {
        "snapshot_schema_version": SNAPSHOT_SCHEMA_VERSION,
        "snapshot_type": "SINGLE_AGENT_CHAT",
        "exported_at": _now_iso(),
        "exported_by": {"user_id": principal["user_id"], "role": principal["role"]},
        "conversation": conversation,
        "messages": messages,
        "production_mutation": False,
        "database_canonical": False,
    }


def _review_snapshot(connection: Any, work_item_id: str, owner: dict[str, str]) -> dict[str, Any]:
    item = _work_item_detail(connection, work_item_id)
    if item["created_by_actor_type"] == "OWNER" and item["created_by_actor_id"] != owner["user_id"]:
        raise HTTPException(status_code=404, detail="Work item not found")
    return {
        "snapshot_schema_version": SNAPSHOT_SCHEMA_VERSION,
        "snapshot_type": "MULTI_AGENT_REVIEW",
        "exported_at": _now_iso(),
        "exported_by": {"user_id": owner["user_id"], "role": owner["role"]},
        "work_item": item,
    }


def _add_meta(document: Document, label: str, value: Any) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}: ")
    run.bold = True
    paragraph.add_run(str(value if value is not None else ""))


def _chat_docx(snapshot: dict[str, Any]) -> bytes:
    document = Document()
    conversation = snapshot["conversation"]
    document.add_heading("Medicine Store Assistant — Conversation Snapshot", level=0)
    _add_meta(document, "Snapshot type", snapshot["snapshot_type"])
    _add_meta(document, "Schema", snapshot["snapshot_schema_version"])
    _add_meta(document, "Exported at", snapshot["exported_at"])
    _add_meta(document, "Conversation ID", conversation.get("conversation_id"))
    _add_meta(document, "Title", conversation.get("title"))
    _add_meta(document, "Agent", f"{conversation.get('agent_display_name', '')} · {conversation.get('agent_call_name', '')}")
    document.add_heading("Transcript", level=1)
    for message in snapshot["messages"]:
        role = "Owner / User" if message.get("role") == "USER" else conversation.get("agent_display_name") or "Assistant"
        document.add_heading(str(role), level=2)
        document.add_paragraph(_clean_human_text(message.get("content")))
        _add_meta(document, "Timestamp", message.get("created_at"))
        provenance = message.get("runtime_provenance") or {}
        if message.get("role") == "ASSISTANT" and provenance:
            _add_meta(
                document,
                "Runtime",
                f"{provenance.get('provider_name') or provenance.get('selected_provider_name') or 'Provider'} · "
                f"{provenance.get('model_name') or provenance.get('selected_model_name') or provenance.get('model_id') or provenance.get('selected_model_id') or 'Model'}"
                + (" · fallback" if provenance.get("fallback_used") else "")
                + (f" · {provenance.get('latency_ms')} ms" if provenance.get("latency_ms") is not None else ""),
            )
        for attachment in message.get("attachments") or []:
            _add_meta(
                document,
                "Attachment",
                f"{attachment.get('filename') or attachment.get('original_filename')} · {attachment.get('content_type')} · {attachment.get('byte_size')} bytes",
            )
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _review_docx(snapshot: dict[str, Any]) -> bytes:
    document = Document()
    item = snapshot["work_item"]
    document.add_heading("Medicine Store Assistant — Multi-Agent Review Snapshot", level=0)
    _add_meta(document, "Schema", snapshot["snapshot_schema_version"])
    _add_meta(document, "Exported at", snapshot["exported_at"])
    _add_meta(document, "Work Item ID", item.get("work_item_id"))
    _add_meta(document, "Title", item.get("title"))
    _add_meta(document, "Status", item.get("status"))
    _add_meta(document, "Production mutation", item.get("production_mutation"))
    _add_meta(document, "Database canonical", item.get("database_canonical"))

    document.add_heading("Conversation / Artifacts", level=1)
    for artifact in item.get("artifacts") or []:
        payload = artifact.get("payload") or {}
        artifact_type = artifact.get("artifact_type")
        if artifact_type == "OWNER_TASK":
            label = "Owner — Task"
            body = payload.get("task")
        elif artifact_type == "OWNER_REVISION":
            label = "Owner — Revision"
            body = payload.get("instruction")
        elif artifact_type == "PARTICIPANT_OUTPUT":
            provenance = payload.get("provenance") or {}
            label = f"{provenance.get('agent_display_name') or payload.get('display_label') or 'Internal agent'} — {payload.get('role') or 'PARTICIPANT'}"
            body = payload.get("response")
        else:
            label = str(artifact_type or "Artifact")
            body = json.dumps(payload, ensure_ascii=False, default=str)
        document.add_heading(label, level=2)
        document.add_paragraph(_clean_human_text(body))
        _add_meta(document, "Artifact", f"v{artifact.get('version')} · {artifact.get('artifact_id')}")
        if artifact_type == "PARTICIPANT_OUTPUT":
            provenance = payload.get("provenance") or {}
            _add_meta(
                document,
                "Runtime",
                f"{provenance.get('selected_provider_name') or 'Provider'} · {provenance.get('selected_model_name') or provenance.get('selected_model_id') or 'Model'}"
                + (" · fallback" if provenance.get("fallback_used") else "")
                + (f" · {provenance.get('latency_ms')} ms" if provenance.get("latency_ms") is not None else ""),
            )

    document.add_heading("Reviews", level=1)
    for review in item.get("reviews") or []:
        document.add_heading(str(review.get("verdict") or "COMMENT"), level=2)
        _add_meta(document, "Reviewed artifact", f"{review.get('artifact_id')} v{review.get('artifact_version')}")
        document.add_paragraph(_clean_human_text(review.get("notes")))

    document.add_heading("Attention", level=1)
    for attention in item.get("attention") or []:
        document.add_paragraph(
            f"{attention.get('category')} · {attention.get('status')} · {attention.get('summary')}"
        )

    document.add_heading("Timeline", level=1)
    for event in item.get("events") or []:
        document.add_paragraph(
            f"{event.get('created_at')} · {event.get('event_type')} · {event.get('actor_type')}"
        )

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


@router.get("/conversations/{conversation_id}/export", summary="Export one owned Single-Agent conversation snapshot")
def export_conversation(
    conversation_id: str,
    format: ExportFormat = Query(default="docx"),
    principal: dict[str, str] = Depends(require_ai_chat_access),
) -> Response:
    engine = _engine()
    try:
        with engine.connect() as connection:
            snapshot = _chat_snapshot(connection, conversation_id, principal)
    finally:
        engine.dispose()
    title = snapshot["conversation"].get("title") or "single-agent-chat"
    if format == "json":
        return _download(_json_bytes(snapshot), media_type="application/json; charset=utf-8", filename=_filename(title, "json"))
    return _download(
        _chat_docx(snapshot),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=_filename(title, "docx"),
    )


@router.get("/multi-agent/work-items/{work_item_id}/export", summary="Export one Owner Multi-Agent Review snapshot")
def export_review(
    work_item_id: str,
    format: ExportFormat = Query(default="docx"),
    owner: dict[str, str] = Depends(require_owner_session),
) -> Response:
    engine = _engine()
    try:
        with engine.connect() as connection:
            snapshot = _review_snapshot(connection, work_item_id, owner)
    finally:
        engine.dispose()
    title = snapshot["work_item"].get("title") or "multi-agent-review"
    if format == "json":
        return _download(_json_bytes(snapshot), media_type="application/json; charset=utf-8", filename=_filename(title, "json"))
    return _download(
        _review_docx(snapshot),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=_filename(title, "docx"),
    )
